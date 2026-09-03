import os
import shutil

from docx import Document


def main() -> None:
    template_path = r"C:\Users\THIS PC\Desktop\Reports\Project_Progress_Report_2026-08-12_ACD.docx"
    out_path = r"C:\Users\THIS PC\Desktop\Reports\Project_Progress_Report_2026-08-12_ACD_TODAY.docx"

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    shutil.copyfile(template_path, out_path)

    doc = Document(out_path)
    if not doc.tables:
        raise RuntimeError("Template docx has no tables to update.")

    tbl = doc.tables[0]

    # Column indices: 0=No., 1=Module/Feature, 2=Description, 3=Status
    module = "ID Photo Upload Size Increase (Up to 50MB)"
    desc = (
        "Increased allowed upload size for ID-related images across registration flows. "
        "Updated validation rules to raise `profile_picture` / `formal_picture` max upload size to "
        "51200KB (about 50MB) in `StudentController`, `EmployeeController`, "
        "`PatronRegistrationController`, `PendingStudentController`, and `PendingEmployeeController`. "
        "Also updated the student create UI hint (`resources/views/students/create.blade.php`) to reflect the new 50MB limit."
    )
    status = "Completed"

    # Task row 1 is table row index 1 (row 0 is the header)
    tbl.cell(1, 0).text = "1"
    tbl.cell(1, 1).text = module
    tbl.cell(1, 2).text = desc
    tbl.cell(1, 3).text = status

    # Clear remaining task rows 2-5 (row indices 2..5)
    for r in [2, 3, 4, 5]:
        for c in [0, 1, 2, 3]:
            tbl.cell(r, c).text = ""

    doc.save(out_path)
    print(f"Wrote report: {out_path}")


if __name__ == "__main__":
    main()

